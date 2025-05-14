import streamlit as st

# QUAN TRỌNG: set_page_config() phải là lệnh Streamlit đầu tiên
st.set_page_config(page_title="Nhập thông tin Khảo sát", layout="wide", page_icon="📋")

# Tiếp theo mới là import các thư viện
import os
import io
import json
import base64
import datetime
import traceback
import uuid
from PIL import Image
import pillow_heif
import requests
from supabase import create_client, Client

# --- Kiểm tra đăng nhập ---
if 'user' not in st.session_state or not st.session_state.user:
    st.warning("Vui lòng đăng nhập để sử dụng App")
    if st.button("Đi đến trang đăng nhập"):
        try:
            st.switch_page("Home.py")
        except Exception:
            try:
                st.switch_page("pages/view_surveys.py")
            except Exception:
                st.error("Không thể chuyển đến trang đăng nhập. Vui lòng quay lại trang chủ.")
    st.stop()

# [Phần còn lại của code giữ nguyên]

# --- Khởi tạo Session State ---
if 'participant_inputs' not in st.session_state:
    st.session_state.participant_inputs = [""]
if 'surveyor_inputs' not in st.session_state:
    st.session_state.surveyor_inputs = [""]
if 'image_uploader_count' not in st.session_state:
    st.session_state.image_uploader_count = 2
if 'uploaded_images' not in st.session_state:
    st.session_state.uploaded_images = {}
if 'survey_data' not in st.session_state:
    st.session_state.survey_data = None
if 'editing_survey_id' not in st.session_state:
    st.session_state.editing_survey_id = None
if 'panel_notes' not in st.session_state:
    st.session_state.panel_notes = []
if 'panel_images' not in st.session_state:
    st.session_state.panel_images = {}

# --- Hàm tiện ích để quản lý người tham gia và người khảo sát ---
def add_participant_input():
    st.session_state.participant_inputs.append("")

def add_surveyor_input():
    st.session_state.surveyor_inputs.append("")

# --- Hàm tiện ích để quản lý upload ảnh ---
def add_image_uploader():
    st.session_state.image_uploader_count += 1

def remove_image_uploader(index):
    if f"image_{index}" in st.session_state.uploaded_images:
        del st.session_state.uploaded_images[f"image_{index}"]
    
    # Reindex các ảnh còn lại
    new_uploaded_images = {}
    new_idx = 0
    for i in range(st.session_state.image_uploader_count):
        if i != index and f"image_{i}" in st.session_state.uploaded_images:
            new_uploaded_images[f"image_{new_idx}"] = st.session_state.uploaded_images[f"image_{i}"]
            new_idx += 1
    
    st.session_state.uploaded_images = new_uploaded_images
    if st.session_state.image_uploader_count > 1:  # Đảm bảo luôn có ít nhất một uploader
        st.session_state.image_uploader_count -= 1

# --- Hàm tiện ích để quản lý Panel Notes ---
def add_panel_note():
    note_id = str(uuid.uuid4())
    st.session_state.panel_notes.append({
        "id": note_id,
        "area": "",
        "device": "",
        "findings": "",
        "images": []
    })
    # Khởi tạo một từ điển với đúng 2 trường uploader ảnh cho note mới
    st.session_state.panel_images[note_id] = {
        "image_0": None,
        "image_1": None
    }

def remove_panel_note(index):
    if index < len(st.session_state.panel_notes):
        note_id = st.session_state.panel_notes[index]["id"]
        # Xóa dữ liệu ảnh của panel note
        if note_id in st.session_state.panel_images:
            del st.session_state.panel_images[note_id]
        # Xóa panel note khỏi danh sách
        st.session_state.panel_notes.pop(index)

def add_panel_image(note_id):
    # Đếm số lượng ảnh hiện có trong panel note
    image_count = len(st.session_state.panel_images.get(note_id, {}))
    # Thêm một slot trống cho ảnh mới
    st.session_state.panel_images.setdefault(note_id, {})[f"image_{image_count}"] = None

def remove_panel_image(note_id, image_key):
    if note_id in st.session_state.panel_images and image_key in st.session_state.panel_images[note_id]:
        del st.session_state.panel_images[note_id][image_key]

# --- Kết nối Supabase ---
def init_supabase():
    """Khởi tạo kết nối Supabase"""
    try:
        if "supabase" not in st.secrets:
            st.error("🔑 Không tìm thấy cấu hình Supabase!")
            return None
            
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        
        if not url or not key:
            st.error("🔑 URL hoặc key Supabase không hợp lệ!")
            return None
            
        # Kết nối
        client = create_client(url, key)
        
        # Kiểm tra kết nối
        try:
            # Thử truy vấn đơn giản để xác minh kết nối
            response = client.table('surveys').select('id').limit(1).execute()
            return client
        except Exception as e:
            st.error(f"❌ Kết nối đến Supabase thất bại: {e}")
            return None
            
    except Exception as e:
        st.error(f"❌ Không thể kết nối đến Supabase: {e}")
        return None

# --- Hàm xử lý ảnh ---
def convert_heic_to_jpeg(file_object):
    """Chuyển đổi ảnh HEIC sang JPEG"""
    try:
        heif_file = pillow_heif.read_heif(file_object.read())
        image = Image.frombytes(
            heif_file.mode,
            heif_file.size,
            heif_file.data,
            "raw"
        )
        output = io.BytesIO()
        image.save(output, format="JPEG")
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"❌ Không thể chuyển đổi .heic -> .jpg: {e}")
        return None

def upload_image_to_supabase(supabase, file_object):
    """Tải ảnh lên Supabase Storage và trả về URL"""
    if not supabase or not file_object:
        return None
        
    try:
        # Sửa tên bucket để khớp với tên hiện có trên Supabase
        bucket_name = "serveyimages"  # Tên bucket đã tồn tại
        
        file_ext = file_object.name.lower().split('.')[-1]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        # Thêm một mã UUID ngẫu nhiên vào tên file để đảm bảo tính duy nhất
        random_id = str(uuid.uuid4())[:8]  # Lấy 8 ký tự đầu của UUID
        unique_filename = f"{timestamp}_{random_id}_{file_object.name.replace(' ', '_')}"
        
        # Xử lý file
        try:
            # Reset con trỏ file
            file_object.seek(0)
            img = Image.open(file_object)
            
            # Xử lý HEIC/HEIF hoặc chuyển đổi RGBA sang RGB
            if file_ext in ['heic', 'heif']:
                file_object.seek(0)
                converted_image = convert_heic_to_jpeg(file_object)
                if not converted_image:
                    return None
                file_content = converted_image.getvalue()
                unique_filename = unique_filename.rsplit('.', 1)[0] + '.jpg'
            elif img.mode == 'RGBA':
                # Chuyển đổi RGBA sang RGB với background trắng
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                # Lưu dưới dạng JPEG
                output = io.BytesIO()
                background.save(output, format='JPEG')
                output.seek(0)
                file_content = output.getvalue()
                unique_filename = unique_filename.rsplit('.', 1)[0] + '.jpg'
            else:
                # Các định dạng khác, sử dụng nguyên bản
                file_object.seek(0)
                file_content = file_object.read()
        except Exception as e:
            # Nếu không thể xử lý ảnh, sử dụng file gốc
            st.warning(f"Không thể xử lý ảnh. Sử dụng file gốc: {e}")
            file_object.seek(0)
            file_content = file_object.read()
        
        # Upload lên Supabase Storage
        supabase.storage.from_(bucket_name).upload(
            path=unique_filename,
            file=file_content,
            file_options={"content-type": f"image/{file_ext if file_ext != 'jpg' else 'jpeg'}"}
        )
        
        # Lấy URL công khai
        file_url = supabase.storage.from_(bucket_name).get_public_url(unique_filename)
        return file_url
        
    except Exception as e:
        st.error(f"Lỗi khi tải ảnh lên: {str(e)}")
        st.error(traceback.format_exc())
        return None

def process_image_for_export(file):
    """Xử lý file ảnh để sử dụng trong export PDF và Word"""
    if file is None:
        return None
    
    file_ext = file.name.lower().split('.')[-1]
    
    try:
        if file_ext in ['heic', 'heif']:
            # Reset con trỏ file
            file.seek(0)
            converted_image = convert_heic_to_jpeg(file)
            if converted_image:
                return Image.open(converted_image)
            return None
        else:
            # Reset con trỏ file
            file.seek(0)
            img = Image.open(file)
            
            # Chuyển đổi RGBA sang RGB nếu cần
            if img.mode == 'RGBA':
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                return background
            
            return img
    except Exception as e:
        st.error(f"Lỗi xử lý ảnh: {e}")
        return None

def load_image_from_url(url):
    """Tải và xử lý ảnh từ URL để sử dụng trong export"""
    try:
        response = requests.get(url)
        if response.status_code == 200:
            img = Image.open(io.BytesIO(response.content))
            
            # Chuyển đổi RGBA sang RGB nếu cần
            if img.mode == 'RGBA':
                # Tạo background trắng
                background = Image.new('RGB', img.size, (255, 255, 255))
                # Paste hình ảnh RGBA lên background
                background.paste(img, mask=img.split()[3])  # Sử dụng kênh alpha làm mask
                return background
            return img
        return None
    except Exception as e:
        st.error(f"Lỗi khi tải ảnh từ URL: {e}")
        return None

# --- Hàm lưu dữ liệu và xuất báo cáo ---
def save_survey_data_to_supabase(supabase, header_data, detail_data, image_urls=None, panel_notes=None, survey_id=None):
    """Lưu dữ liệu khảo sát vào Supabase"""
    if not supabase:
        return None
        
    try:
        # Tạo đối tượng dữ liệu khảo sát
        survey_data = {
            "company_name": header_data["company_name"],
            "address": header_data["address"],
            "phone": header_data["phone"],
            "survey_date": header_data["survey_date"],
            "participants": header_data["participants"],
            "surveyors": header_data["surveyors"],
            "created_by": st.session_state.user["id"],
            "area": detail_data[0],
            "device": detail_data[1],
            "findings": detail_data[2],
            "images": image_urls if image_urls else []
        }
        
        # Insert hoặc Update dữ liệu trong Supabase
        if survey_id:
            # Kiểm tra quyền nếu không phải admin
            if st.session_state.user["role"] != "admin":
                response = supabase.table('surveys').select('created_by').eq('id', survey_id).execute()
                if not response.data or response.data[0]["created_by"] != st.session_state.user["id"]:
                    st.error("Bạn không có quyền chỉnh sửa khảo sát này!")
                    return None
            
            # Cập nhật khảo sát hiện có
            response = supabase.table('surveys').update(survey_data).eq('id', survey_id).execute()
            st.success(f"✅ Khảo sát đã được cập nhật thành công!")
            
            # Xóa các panel notes cũ
            supabase.table('panel_notes').delete().eq('survey_id', survey_id).execute()
        else:
            # Tạo khảo sát mới
            response = supabase.table('surveys').insert(survey_data).execute()
        
        if response.data:
            survey_id = response.data[0].get('id')
            
            # Lưu các panel notes
            if panel_notes:
                for note in panel_notes:
                    note_data = {
                        "survey_id": survey_id,
                        "area": note["area"],
                        "device": note["device"],
                        "findings": note["findings"],
                        "images": note["images"],
                        "created_by": st.session_state.user["id"]
                    }
                    supabase.table('panel_notes').insert(note_data).execute()
            
            return survey_id
        else:
            st.error("Không nhận được dữ liệu phản hồi từ Supabase")
            return None
            
    except Exception as e:
        st.error(f"Lỗi khi lưu dữ liệu vào Supabase: {str(e)}")
        st.error(traceback.format_exc())
        return None

def delete_survey_from_supabase(supabase, survey_id):
    """Xóa khảo sát từ Supabase"""
    if not supabase or not survey_id:
        return False
        
    try:
        # Kiểm tra quyền nếu không phải admin
        if st.session_state.user["role"] != "admin":
            response = supabase.table('surveys').select('created_by').eq('id', survey_id).execute()
            if not response.data or response.data[0]["created_by"] != st.session_state.user["id"]:
                st.error("Bạn không có quyền xóa khảo sát này!")
                return False
        
        # Xóa khảo sát (các panel_notes sẽ tự động bị xóa do có constraint ON DELETE CASCADE)
        response = supabase.table('surveys').delete().eq('id', survey_id).execute()
        return True
    except Exception as e:
        st.error(f"Lỗi khi xóa khảo sát: {str(e)}")
        return False

def get_download_link(file_content, file_name, display_text):
    """Tạo link tải xuống cho các file được tạo."""
    b64 = base64.b64encode(file_content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}" class="download-button">{display_text}</a>'
    return href

def export_to_pdf(survey_data, images, panel_notes=None):
    """Tạo file PDF từ dữ liệu khảo sát."""
    # Import các module cần thiết
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os
    
    buffer = io.BytesIO()
    
    # Đăng ký font hỗ trợ tiếng Việt
    try:
        # Thiết lập các đường dẫn font có thể có
        current_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
        
        # Danh sách font để thử
        font_files = [
            ('TimesTiengViet', 'times.ttf', 'Times New Roman'),
            ('Arial', 'arial.ttf', 'Arial'),
            ('DejaVuSans', 'DejaVuSans.ttf', 'DejaVuSans'),
        ]
        
        # Các đường dẫn có thể chứa font
        possible_paths = [
            os.path.join(current_dir, 'assets', 'fonts'),
            os.path.join(current_dir, 'assets'),
            current_dir,
            '/usr/share/fonts/truetype',
            '/usr/share/fonts/truetype/dejavu',
            '/usr/share/fonts/TTF',
            'C:\\Windows\\Fonts',  # Đường dẫn Windows
        ]
        
        # Thử đăng ký từng font
        font_registered = False
        font_name = 'Times-Roman'  # Font mặc định của ReportLab
        
        for font_key, font_file, display_name in font_files:
            if font_registered:
                break
                
            for path in possible_paths:
                try:
                    full_path = os.path.join(path, font_file)
                    if os.path.exists(full_path):
                        pdfmetrics.registerFont(TTFont(font_key, full_path))
                        font_name = font_key
                        font_registered = True
                        st.success(f"Đã đăng ký font {display_name} từ {full_path}")
                        break
                except Exception as e:
                    continue
        
        # Nếu không đăng ký được font tùy chỉnh, sử dụng font có sẵn trong ReportLab
        if not font_registered:
            font_name = 'Times-Roman'  # Times-Roman là font có sẵn trong ReportLab, tương tự Times New Roman
            st.info(f"Sử dụng font mặc định: {font_name}")
    except Exception as e:
        st.write(f"Lỗi khi đăng ký font: {e}")
        font_name = 'Times-Roman'  # Mặc định ở đây cũng là Times-Roman
    
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Tạo style cho tiêu đề và nội dung
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=16,
        alignment=1,
        spaceAfter=12
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        spaceAfter=6
    )
    
    content = []
    
    # Tiêu đề báo cáo
    content.append(Paragraph("BÁO CÁO KHẢO SÁT ISO 50001:2018", title_style))
    content.append(Spacer(1, 20))
    
    # Thông tin công ty và khảo sát
    header_data = [
        ["Tên công ty:", survey_data['header']['company_name']],
        ["Địa chỉ:", survey_data['header']['address']],
        ["Số điện thoại:", survey_data['header']['phone']],
        ["Ngày khảo sát:", survey_data['header']['survey_date']],
    ]
    
    t = Table(header_data, colWidths=[120, 300])
    t.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), font_name, 10),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    content.append(t)
    content.append(Spacer(1, 10))
    
    # Người tham gia
    content.append(Paragraph("Thành viên tham gia:", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name)))
    participants = survey_data['header']['participants'].split('\n')
    for p in participants:
        if p.strip():
            content.append(Paragraph(f"• {p}", normal_style))
    content.append(Spacer(1, 10))
    
    # Người khảo sát
    content.append(Paragraph("Người khảo sát:", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name)))
    surveyors = survey_data['header']['surveyors'].split('\n')
    for s in surveyors:
        if s.strip():
            content.append(Paragraph(f"• {s}", normal_style))
    content.append(Spacer(1, 20))
    
    # Panel Notes - Hiển thị ngay từ panel note đầu tiên
    if panel_notes and len(panel_notes) > 0:
        # Bắt đầu với Panel Note đầu tiên thay vì CHI TIẾT KHẢO SÁT
        for note_idx, note in enumerate(panel_notes):
            if note_idx == 0:
                content.append(Paragraph(f"KHU VỰC KHẢO SÁT #{note_idx+1}", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
            else:
                content.append(Paragraph(f"KHU VỰC KHẢO SÁT #{note_idx+1}", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
            
            content.append(Spacer(1, 10))
            
            note_data = [
                ["Khu vực:", note['area']],
                ["Thiết bị:", note['device']],
                ["Mô tả tổn thất/thông số kỹ thuật:", note['findings']],
            ]
            
            t = Table(note_data, colWidths=[150, 270])
            t.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            content.append(t)
            content.append(Spacer(1, 10))
            
            # Hình ảnh của panel note
            if note.get('images') and len(note['images']) > 0:
                content.append(Paragraph(f"ẢNH KHẢO SÁT KHU VỰC #{note_idx+1}", ParagraphStyle('Heading3', parent=styles['Heading3'], fontName=font_name, alignment=1)))
                content.append(Spacer(1, 10))
                
                for img_idx, img_url in enumerate(note['images']):
                    img = load_image_from_url(img_url)
                    if img:
                        # Chuyển đổi sang RGB nếu là RGBA
                        if img.mode == 'RGBA':
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[3])
                            img = background
                        
                        # Chuyển đổi PIL Image sang bytesIO để ReportLab có thể sử dụng
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG')
                        img_byte_arr.seek(0)
                        
                        # Thêm caption cho hình ảnh
                        content.append(Paragraph(f"Hình {note_idx+1}.{img_idx+1}:", normal_style))
                        
                        # Tính toán kích thước ảnh phù hợp
                        img_width, img_height = img.size
                        ratio = min(400/img_width, 300/img_height) if img_width > 0 and img_height > 0 else 1
                        new_width = img_width * ratio
                        new_height = img_height * ratio
                        
                        # Thêm hình ảnh vào tài liệu PDF
                        img_reportlab = RLImage(img_byte_arr, width=new_width, height=new_height)
                        content.append(img_reportlab)
                        content.append(Spacer(1, 10))
            
            content.append(Spacer(1, 10))
    else:
        # Nếu không có panel notes, hiển thị thông tin từ detail_data
        content.append(Paragraph("KHU VỰC KHẢO SÁT #1", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
        content.append(Spacer(1, 10))
        
        detail_data = [
            ["Khu vực:", survey_data['detail'][0]],
            ["Thiết bị:", survey_data['detail'][1]],
            ["Mô tả tổn thất/thông số kỹ thuật:", survey_data['detail'][2]],
        ]
        
        t = Table(detail_data, colWidths=[150, 270])
        t.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), font_name, 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        content.append(t)
        content.append(Spacer(1, 20))
        
        # Hình ảnh
        if images:
            content.append(Paragraph("ẢNH KHẢO SÁT KHU VỰC #1", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
            content.append(Spacer(1, 10))
            
            for i, img in enumerate(images):
                if img:
                    # Chuyển đổi sang RGB nếu là RGBA
                    if img.mode == 'RGBA':
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[3])
                        img = background
                    
                    # Chuyển đổi PIL Image sang bytesIO để ReportLab có thể sử dụng
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='JPEG')
                    img_byte_arr.seek(0)
                    
                    # Thêm caption cho hình ảnh
                    content.append(Paragraph(f"Hình 1.{i+1}:", normal_style))
                    
                    # Tính toán kích thước ảnh phù hợp (tối đa 400x300, giữ tỷ lệ)
                    img_width, img_height = img.size
                    ratio = min(400/img_width, 300/img_height) if img_width > 0 and img_height > 0 else 1
                    new_width = img_width * ratio
                    new_height = img_height * ratio
                    
                    # Thêm hình ảnh vào tài liệu PDF
                    img_reportlab = RLImage(img_byte_arr, width=new_width, height=new_height)
                    content.append(img_reportlab)
                    content.append(Spacer(1, 10))
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    content.append(Spacer(1, 20))
    content.append(Paragraph(f"Báo cáo được xuất ngày: {current_date}", normal_style))
    content.append(Paragraph(f"Người xuất báo cáo: {st.session_state.user.get('full_name', '')}", normal_style))
    
    doc.build(content)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def export_to_word(survey_data, images, panel_notes=None):
    """Tạo file Word từ dữ liệu khảo sát."""
    # Import các module cần thiết
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Thiết lập font và cỡ chữ mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Tiêu đề
    heading = doc.add_heading('BÁO CÁO KHẢO SÁT ISO 50001', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin công ty và khảo sát
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    rows = table.rows
    cells = rows[0].cells
    cells[0].text = 'Tên công ty:'
    cells[1].text = survey_data['header']['company_name']
    
    cells = rows[1].cells
    cells[0].text = 'Địa chỉ:'
    cells[1].text = survey_data['header']['address']
    
    cells = rows[2].cells
    cells[0].text = 'Số điện thoại:'
    cells[1].text = survey_data['header']['phone']
    
    cells = rows[3].cells
    cells[0].text = 'Ngày khảo sát:'
    cells[1].text = survey_data['header']['survey_date']
    
    doc.add_paragraph('')
    
    # Người tham gia
    doc.add_heading('Thành viên tham gia:', level=2)
    participants = survey_data['header']['participants'].split('\n')
    for p in participants:
        if p.strip():
            doc.add_paragraph(f"• {p}", style='List Bullet')
    
    # Người khảo sát
    doc.add_heading('Người khảo sát:', level=2)
    surveyors = survey_data['header']['surveyors'].split('\n')
    for s in surveyors:
        if s.strip():
            doc.add_paragraph(f"• {s}", style='List Bullet')
    
    doc.add_paragraph('')
    
    # Panel Notes - Hiển thị ngay từ panel note đầu tiên
    if panel_notes and len(panel_notes) > 0:
        # Bắt đầu với Panel Note đầu tiên thay vì CHI TIẾT KHẢO SÁT
        for note_idx, note in enumerate(panel_notes):
            heading = doc.add_heading(f'KHU VỰC KHẢO SÁT #{note_idx+1}', level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            
            rows = table.rows
            cells = rows[0].cells
            cells[0].text = 'Khu vực:'
            cells[1].text = note['area']
            
            cells = rows[1].cells
            cells[0].text = 'Thiết bị:'
            cells[1].text = note['device']
            
            cells = rows[2].cells
            cells[0].text = 'Mô tả tổn thất/thông số kỹ thuật:'
            cells[1].text = note['findings']
            
            doc.add_paragraph('')
            
            # Hình ảnh của panel note
            if note.get('images') and len(note['images']) > 0:
                heading = doc.add_heading(f'ẢNH KHẢO SÁT KHU VỰC #{note_idx+1}', level=3)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for img_idx, img_url in enumerate(note['images']):
                    img = load_image_from_url(img_url)
                    if img:
                        # Chuyển đổi sang RGB nếu là RGBA
                        if img.mode == 'RGBA':
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[3])
                            img = background
                        
                        # Lưu ảnh tạm thời để chèn vào Word
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG')
                        img_byte_arr.seek(0)
                        
                        # Thêm caption và hình ảnh
                        doc.add_paragraph(f"Hình {note_idx+1}.{img_idx+1}:")
                        doc.add_picture(img_byte_arr, width=Inches(6))
                        doc.add_paragraph('')
            
            doc.add_paragraph('')
    else:
        # Nếu không có panel notes, hiển thị thông tin từ detail_data
        heading = doc.add_heading('KHU VỰC KHẢO SÁT #1', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        rows = table.rows
        cells = rows[0].cells
        cells[0].text = 'Khu vực:'
        cells[1].text = survey_data['detail'][0]
        
        cells = rows[1].cells
        cells[0].text = 'Thiết bị:'
        cells[1].text = survey_data['detail'][1]
        
        cells = rows[2].cells
        cells[0].text = 'Mô tả tổn thất/thông số kỹ thuật:'
        cells[1].text = survey_data['detail'][2]
        
        doc.add_paragraph('')
        
        # Hình ảnh
        if images:
            heading = doc.add_heading('ẢNH KHẢO SÁT KHU VỰC #1', level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, img in enumerate(images):
                if img:
                    # Chuyển đổi sang RGB nếu là RGBA
                    if img.mode == 'RGBA':
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[3])
                        img = background
                    
                    # Lưu ảnh tạm thời để chèn vào Word
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='JPEG')
                    img_byte_arr.seek(0)
                    
                    # Thêm caption và hình ảnh
                    doc.add_paragraph(f"Hình 1.{i+1}:")
                    doc.add_picture(img_byte_arr, width=Inches(6))
                    doc.add_paragraph('')
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    doc.add_paragraph('')
    doc.add_paragraph(f"Báo cáo được xuất ngày: {current_date}")
    doc.add_paragraph(f"Người xuất báo cáo: {st.session_state.user.get('full_name', '')}")
    
    # Lưu vào memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    docx = buffer.getvalue()
    buffer.close()
    return docx

# --- CSS cho giao diện ---
st.markdown("""
<style>
.download-button {
    display: inline-block;
    padding: 0.5em 1em;
    text-decoration: none;
    color: white;
    background-color: #0066cc;
    border-radius: 5px;
    font-weight: bold;
    margin: 0.5em 0;
    text-align: center;
}
.download-button:hover {
    background-color: #0052a3;
}
.image-controls {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin: 10px 0;
}
.stButton button {
    min-height: 2.5em;
}
.image-preview {
    margin: 10px 0;
    border: 1px solid #ddd;
    border-radius: 5px;
    padding: 10px;
}
.panel-note {
    margin: 15px 0;
    padding: 15px;
    border: 1px solid #ddd;
    border-radius: 5px;
    background-color: #f9f9f9;
}
.user-info {
    text-align: right;
    font-size: 0.9em;
    margin-bottom: 10px;
}
</style>
""", unsafe_allow_html=True)

# --- Hiển thị thông tin người dùng ---
user_role = "Quản trị viên" if st.session_state.user["role"] == "admin" else "Thành viên"
st.markdown(f"""
<div class="user-info">
    Xin chào, <b>{st.session_state.user['full_name']}</b> | Vai trò: <b>{user_role}</b> | 
    <a href="javascript:void(0);" id="logout-link">Đăng xuất</a>
</div>
<script>
    document.getElementById('logout-link').addEventListener('click', function() {{
        window.parent.postMessage({{type: 'streamlit:setComponentValue', value: true, dataType: 'logout'}}, '*');
    }});
</script>
""", unsafe_allow_html=True)

# Xử lý đăng xuất
if st.session_state.get('logout', False):
    st.session_state.user = None
    st.session_state.logout = False
    try:
        st.switch_page("login.py")
    except Exception:
        try:
            st.switch_page("../login.py")
        except Exception:
            st.error("Không thể chuyển đến trang đăng nhập. Vui lòng quay lại trang chủ.")
    st.stop()

# --- Kết nối với Supabase ---
supabase = init_supabase()

# --- Giao diện chính ---
st.title("📝 Form Khảo sát hiện trường ISO 50001:2018")

# Tab cho các chức năng khác nhau
tab_input, tab_export = st.tabs(["Nhập liệu", "Xuất báo cáo"])

with tab_input:
    if not supabase:
        st.error("😢 Không thể kết nối đến cơ sở dữ liệu Supabase.")
        st.warning("Vui lòng cấu hình kết nối Supabase trong Streamlit Secrets.")
        st.info("""
        ### Hướng dẫn cấu hình Supabase
        
        1. Đăng nhập vào [Streamlit Cloud](https://share.streamlit.io)
        2. Chọn Settings > Secrets của ứng dụng
        3. Thêm thông tin sau (với giá trị thực tế của bạn):
        ```
        [supabase]
        url = "https://your-project-id.supabase.co"
        key = "your-supabase-service-role-key"
        ```
        4. Lưu và khởi động lại ứng dụng
        """)
        
        # Hiển thị giao diện giả để người dùng vẫn có thể xem giao diện
        st.subheader("Xem trước giao diện (chế độ demo)")
        st.write("Lưu ý: Chức năng lưu dữ liệu sẽ không hoạt động cho đến khi cấu hình Supabase")
        
        with st.form("demo_form"):
            st.header("Thông tin khảo sát")
            st.text_input("Tên công ty")
            st.text_input("Địa chỉ")
            st.text_input("Số điện thoại")
            st.date_input("Ngày khảo sát")
            st.form_submit_button("Gửi dữ liệu (Chức năng demo)")
        
        st.stop()  # Dừng ở đây, không chạy phần code thực tế
    
    # Form nhập dữ liệu - KHÔNG CHỨa nút Lưu dữ liệu nữa
    with st.form("survey_form", clear_on_submit=False):
        st.header("Thông tin khảo sát ISO 50001")
        
        # Nếu đang chỉnh sửa, hiển thị ID
        if st.session_state.editing_survey_id:
            st.info(f"Đang chỉnh sửa khảo sát ID: {st.session_state.editing_survey_id}")
            
            # Load dữ liệu hiện có nếu đang ở chế độ chỉnh sửa
            try:
                response = supabase.table('surveys').select('*').eq('id', st.session_state.editing_survey_id).execute()
                if response.data and len(response.data) > 0:
                    survey_data = response.data[0]
                    
                    # Kiểm tra quyền nếu không phải admin
                    if st.session_state.user["role"] != "admin" and survey_data["created_by"] != st.session_state.user["id"]:
                        st.error("Bạn không có quyền chỉnh sửa khảo sát này!")
                        st.button("Quay lại")
                        st.stop()
                    
                    # Load panel notes
                    panel_notes_response = supabase.table('panel_notes').select('*').eq('survey_id', st.session_state.editing_survey_id).execute()
                    if panel_notes_response.data:
                        st.session_state.panel_notes = panel_notes_response.data
                        
                        # Khởi tạo panel_images
                        for note in st.session_state.panel_notes:
                            note_id = note["id"]
                            st.session_state.panel_images[note_id] = {}
            except Exception as e:
                st.error(f"Lỗi khi tải dữ liệu: {e}")
        
        company_name = st.text_input("Tên công ty")
        address = st.text_input("Địa chỉ")
        phone = st.text_input("Số điện thoại")
        survey_date = st.date_input("Ngày khảo sát")

        # Thành viên tham gia
        st.subheader("Thành viên tham gia")
        participant_values = []
        for i in range(len(st.session_state.participant_inputs)):
            participant_values.append(st.text_input(f"Thành viên {i+1}", key=f"participant_{i}"))
        
        if st.form_submit_button("➕ Thêm thành viên", type="secondary"):
            add_participant_input()
            st.rerun()

        # Người khảo sát
        st.subheader("Người khảo sát")
        surveyor_values = []
        for i in range(len(st.session_state.surveyor_inputs)):
            surveyor_values.append(st.text_input(f"Người khảo sát {i+1}", key=f"surveyor_{i}"))
        
        if st.form_submit_button("➕ Thêm người khảo sát", type="secondary"):
            add_surveyor_input()
            st.rerun()

        # Thông tin thiết bị
        st.subheader("Thông tin thiết bị")
        area = st.text_input("Khu vực")
        device = st.text_input("Thiết bị")
        findings = st.text_area("Mô tả tổn thất hoặc thông số kỹ thuật")
        
        # Bỏ nút submit form ở đây
    
    # Phần upload ảnh (đặt ngoài form để tránh lỗi)
    st.subheader("Tải lên hình ảnh chính")
    
    for i in range(st.session_state.image_uploader_count):
        cols = st.columns([4, 1])
        with cols[0]:
            uploaded_file = st.file_uploader(
                f"Ảnh {i+1}", 
                type=["png", "jpg", "jpeg", "heic", "heif", "bmp"], 
                key=f"image_{i}"
            )
            if uploaded_file:
                st.session_state.uploaded_images[f"image_{i}"] = uploaded_file
                try:
                    image = Image.open(uploaded_file)
                    st.image(image, caption=f"Ảnh {i+1}: {uploaded_file.name}", width=300)
                except Exception as e:
                    st.error(f"Không thể hiển thị ảnh: {e}")
        
        with cols[1]:
            if st.button(f"Xóa ảnh #{i+1}", key=f"remove_img_{i}"):
                remove_image_uploader(i)
                st.rerun()
    
    # Nút thêm ảnh (đặt ngoài form)
    if st.button("➕ Thêm ảnh chính khác"):
        add_image_uploader()
        st.rerun()
    
    # Panel Notes
    st.header("Khảo sát Khu vực/Quá trình/Thiết bị")
    st.write("Thêm các ghi chú chi tiết cho từng khu vực, Quá trình hoặc thiết bị")
    
    # Hiển thị các panel notes hiện có
    for idx, note in enumerate(st.session_state.panel_notes):
        with st.expander(f"Khu vực khảo sát #{idx+1}", expanded=True):
            st.markdown(f"<div class='panel-note'>", unsafe_allow_html=True)
            
            # Form nhập thông tin panel note
            note_cols = st.columns(3)
            with note_cols[0]:
                note["area"] = st.text_input(f"Khu vực", value=note.get("area", ""), key=f"note_{idx}_area")
            with note_cols[1]:
                note["device"] = st.text_input(f"Thiết bị", value=note.get("device", ""), key=f"note_{idx}_device")
            with note_cols[2]:
                if st.button(f"Xóa Khu vực khảo sát", key=f"remove_note_{idx}"):
                    remove_panel_note(idx)
                    st.rerun()
            
            note["findings"] = st.text_area(f"Mô tả tổn thất hoặc thông số kỹ thuật", value=note.get("findings", ""), key=f"note_{idx}_findings")
            
            # Upload hình ảnh cho panel note
            st.subheader(f"Hình ảnh Khu vực khảo sát #{idx+1}")
            
            note_id = note["id"]
            # Hiển thị ảnh đã có từ CSDL nếu đang chỉnh sửa
            existing_images = note.get("images", [])
            if existing_images:
                image_cols = st.columns(min(3, len(existing_images)))
                for img_idx, img_url in enumerate(existing_images):
                    with image_cols[img_idx % 3]:
                        st.image(img_url, caption=f"Ảnh {img_idx+1}", width=200)
                        if st.button(f"Xóa ảnh này", key=f"remove_note_img_{idx}_{img_idx}"):
                            existing_images.pop(img_idx)
                            st.rerun()
            
            # Thêm ảnh mới - chỉ hiển thị số lượng trường uploader đã định nghĩa
            panel_images = st.session_state.panel_images.get(note_id, {})
            
            for img_idx in range(len(panel_images)):
                img_key = f"image_{img_idx}"
                if img_key in panel_images:
                    img_cols = st.columns([4, 1])
                    with img_cols[0]:
                        uploaded_file = st.file_uploader(
                            f"Thêm ảnh {img_idx+1}", 
                            type=["png", "jpg", "jpeg", "heic", "heif", "bmp"], 
                            key=f"note_{idx}_img_{img_idx}"
                        )
                        if uploaded_file:
                            panel_images[img_key] = uploaded_file
                            try:
                                image = Image.open(uploaded_file)
                                st.image(image, caption=f"Ảnh {img_idx+1}: {uploaded_file.name}", width=300)
                            except Exception as e:
                                st.error(f"Không thể hiển thị ảnh: {e}")
                    
                    with img_cols[1]:
                        if st.button(f"Xóa", key=f"remove_note_img_new_{idx}_{img_idx}"):
                            remove_panel_image(note_id, img_key)
                            st.rerun()
            
            if st.button(f"➕ Thêm ảnh khác", key=f"add_note_img_{idx}"):
                add_panel_image(note_id)
                st.rerun()
            
            st.markdown(f"</div>", unsafe_allow_html=True)
    
    # Di chuyển nút Thêm Panel Note và Lưu dữ liệu xuống cuối trang
    col1, col2 = st.columns(2)
    with col1:
        if st.button("➕ Thêm Khu vực khảo sát"):
            add_panel_note()
            st.rerun()
    
    with col2:
        # Nút Lưu dữ liệu được đưa ra khỏi form và đặt ở cuối
        submitted = st.button("Lưu dữ liệu", type="primary")
    
    # Xử lý nút submit
    if submitted:
        with st.spinner("Đang xử lý dữ liệu..."):
            # Lấy dữ liệu từ form
            participants = "\n".join([v for v in participant_values if v])
            surveyors = "\n".join([v for v in surveyor_values if v])
            
            header_data = {
                "company_name": company_name,
                "address": address,
                "phone": phone,
                "survey_date": str(survey_date),
                "participants": participants,
                "surveyors": surveyors
            }

            detail_data = [area, device, findings]
            
            # Upload ảnh lên Supabase
            image_urls = []
            
            for i in range(st.session_state.image_uploader_count):
                file_key = f"image_{i}"
                if file_key in st.session_state.uploaded_images:
                    file = st.session_state.uploaded_images[file_key]
                    if file:
                        with st.spinner(f"Đang tải lên ảnh {i+1}..."):
                            image_url = upload_image_to_supabase(supabase, file)
                            if image_url:
                                image_urls.append(image_url)
                                st.success(f"✅ Ảnh {i+1} đã được tải lên")
                            else:
                                st.error(f"❌ Không thể tải lên ảnh {i+1}")
            
            # Xử lý ảnh cho panel notes
            panel_notes_data = []
            for idx, note in enumerate(st.session_state.panel_notes):
                note_id = note["id"]
                note_data = {
                    "id": note_id,
                    "area": note["area"],
                    "device": note["device"],
                    "findings": note["findings"],
                    "images": note.get("images", [])  # Giữ lại ảnh cũ nếu có
                }
                
                # Upload ảnh mới cho panel note
                if note_id in st.session_state.panel_images:
                    for img_key, file in st.session_state.panel_images[note_id].items():
                        if file:
                            with st.spinner(f"Đang tải lên ảnh cho Khu vực khảo sát #{idx+1}..."):
                                image_url = upload_image_to_supabase(supabase, file)
                                if image_url:
                                    note_data["images"].append(image_url)
                                    st.success(f"✅ Ảnh cho Khu vực khảo sát #{idx+1} đã được tải lên")
                                else:
                                    st.error(f"❌ Không thể tải lên ảnh cho Khu vực khảo sát #{idx+1}")
                
                panel_notes_data.append(note_data)

            # Lưu dữ liệu vào Supabase
            with st.spinner("Đang lưu dữ liệu khảo sát..."):
                survey_id = save_survey_data_to_supabase(
                    supabase, 
                    header_data, 
                    detail_data, 
                    image_urls,
                    panel_notes_data,
                    st.session_state.editing_survey_id
                )
                
                if survey_id:
                    action = "cập nhật" if st.session_state.editing_survey_id else "lưu"
                    st.success(f"🎉 Dữ liệu khảo sát đã được {action} thành công! (ID: {survey_id})")
                    
                    # Lưu dữ liệu vào session state để sử dụng ở tab xuất báo cáo
                    st.session_state.survey_data = {
                        "header": header_data,
                        "detail": detail_data,
                        "image_urls": image_urls,
                        "panel_notes": panel_notes_data
                    }
                    
                    # Reset trạng thái chỉnh sửa
                    st.session_state.editing_survey_id = None
                    
                    st.info("Bạn có thể chuyển sang tab 'Xuất báo cáo' để tải báo cáo dưới dạng PDF hoặc Word.")
                else:
                    st.error("❌ Lưu dữ liệu thất bại.")

with tab_export:
    st.header("Xuất báo cáo khảo sát ISO 50001")
    
    # Kiểm tra xem có dữ liệu để xuất không
    if st.session_state.survey_data:
        st.info("Bạn có thể tải xuống báo cáo khảo sát dưới dạng PDF hoặc Word.")
        
        # Xử lý hình ảnh để sẵn sàng xuất báo cáo
        images = []
        
        if "image_urls" in st.session_state.survey_data:
            # Tải ảnh từ URL (Supabase)
            with st.spinner("Đang tải ảnh để chuẩn bị xuất báo cáo..."):
                for url in st.session_state.survey_data["image_urls"]:
                    img = load_image_from_url(url)
                    if img:
                        images.append(img)
        else:
            # Sử dụng ảnh upload local (cách cũ)
            for i in range(st.session_state.image_uploader_count):
                if f"image_{i}" in st.session_state.uploaded_images:
                    file = st.session_state.uploaded_images[f"image_{i}"]
                    if file:
                        # Lưu lại vị trí con trỏ đọc file
                        current_pos = file.tell()
                        # Đặt lại vị trí con trỏ file về đầu để đọc
                        file.seek(0)
                        img = process_image_for_export(file)
                        # Khôi phục vị trí con trỏ ban đầu
                        file.seek(current_pos)
                        images.append(img)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Tải xuống Báo cáo khảo sát pdf")
            if st.button("Tạo file PDF"):
                with st.spinner("Đang tạo file PDF..."):
                    pdf_data = export_to_pdf(
                        st.session_state.survey_data, 
                        images, 
                        st.session_state.survey_data.get("panel_notes", [])
                    )
                    company_name_safe = st.session_state.survey_data['header']['company_name'].replace(' ', '_')
                    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"bao_cao_khao_sat_{company_name_safe}_{date_str}.pdf"
                    
                    st.markdown(
                        get_download_link(pdf_data, filename, "📥 Tải xuống Báo cáo file PDF"),
                        unsafe_allow_html=True
                    )
        
        with col2:
            st.subheader("Tải xuống Báo cáo khảo sát file Word")
            if st.button("Tạo file Báo cáo khảo sát Word"):
                with st.spinner("Đang tạo file Báo cáo Word..."):
                    docx_data = export_to_word(
                        st.session_state.survey_data, 
                        images,
                        st.session_state.survey_data.get("panel_notes", [])
                    )
                    company_name_safe = st.session_state.survey_data['header']['company_name'].replace(' ', '_')
                    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"bao_cao_khao_sat_{company_name_safe}_{date_str}.docx"
                    
                    st.markdown(
                        get_download_link(docx_data, filename, "📥 Tải xuống file Word"),
                        unsafe_allow_html=True
                    )
    else:
        st.warning("Chưa có dữ liệu để xuất báo cáo. Vui lòng nhập và gửi dữ liệu khảo sát trước.")