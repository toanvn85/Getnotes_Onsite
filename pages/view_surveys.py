import streamlit as st

# QUAN TRỌNG: set_page_config() phải là lệnh Streamlit đầu tiên
st.set_page_config(page_title="Xem danh sách khảo sát", layout="wide", page_icon="📋")

# Tiếp theo mới là import các thư viện
import io
import base64
import datetime
import traceback
import os
from PIL import Image
import requests
from supabase import create_client

# --- Kiểm tra đăng nhập ---
if 'user' not in st.session_state or not st.session_state.user:
    st.warning("Vui lòng đăng nhập để sử dụng ứng dụng")
    if st.button("Đi đến trang đăng nhập"):
        try:
            st.switch_page("../Home.py")
        except Exception:
            try:
                st.switch_page("Getnotes_Onsite.py")
            except Exception:
                st.error("Không thể chuyển đến trang đăng nhập. Vui lòng quay lại trang chủ.")
    st.stop()

# [Phần còn lại của code giữ nguyên]

# --- CSS cho giao diện ---
st.markdown("""
<style>
.download-button {
    display: inline-block;
    padding: 0.5em 1em;
    text-decoration: none;
    color: white;
    background-color: #0066cc;
    border-radius: 5px;
    font-weight: bold;
    margin: 0.5em 0;
    text-align: center;
}
.download-button:hover {
    background-color: #0052a3;
}
.stButton button {
    min-height: 2.5em;
}
.user-info {
    text-align: right;
    font-size: 0.9em;
    margin-bottom: 10px;
}
</style>
""", unsafe_allow_html=True)

# --- Hiển thị thông tin người dùng ---
user_role = "Quản trị viên" if st.session_state.user["role"] == "admin" else "Thành viên"
st.markdown(f"""
<div class="user-info">
    Xin chào, <b>{st.session_state.user['full_name']}</b> | Vai trò: <b>{user_role}</b> | 
    <a href="javascript:void(0);" id="logout-link">Đăng xuất</a>
</div>
<script>
    document.getElementById('logout-link').addEventListener('click', function() {{
        window.parent.postMessage({{type: 'streamlit:setComponentValue', value: true, dataType: 'logout'}}, '*');
    }});
</script>
""", unsafe_allow_html=True)

# Xử lý đăng xuất
if st.session_state.get('logout', False):
    st.session_state.user = None
    st.session_state.logout = False
    try:
        st.switch_page("../login.py")
    except Exception:
        try:
            st.switch_page("login.py")
        except Exception:
            st.error("Không thể chuyển đến trang đăng nhập. Vui lòng quay lại trang chủ.")
    st.stop()

# --- Khởi tạo Session State ---
if 'editing_survey_id' not in st.session_state:
    st.session_state.editing_survey_id = None

# --- Các hàm tiện ích ---
def init_supabase():
    """Khởi tạo kết nối Supabase"""
    try:
        if "supabase" not in st.secrets:
            st.error("🔑 Không tìm thấy cấu hình Supabase!")
            return None
            
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        
        if not url or not key:
            st.error("🔑 URL hoặc key Supabase không hợp lệ!")
            return None
            
        # Kết nối
        client = create_client(url, key)
        
        # Kiểm tra kết nối
        try:
            # Thử truy vấn đơn giản để xác minh kết nối
            response = client.table('surveys').select('id').limit(1).execute()
            return client
        except Exception as e:
            st.error(f"❌ Kết nối đến Supabase thất bại: {e}")
            return None
            
    except Exception as e:
        st.error(f"❌ Không thể kết nối đến Supabase: {e}")
        return None

def load_image_from_url(url):
    """Tải và xử lý ảnh từ URL để sử dụng trong export"""
    try:
        response = requests.get(url)
        if response.status_code == 200:
            img = Image.open(io.BytesIO(response.content))
            
            # Chuyển đổi RGBA sang RGB nếu cần
            if img.mode == 'RGBA':
                # Tạo background trắng
                background = Image.new('RGB', img.size, (255, 255, 255))
                # Paste hình ảnh RGBA lên background
                background.paste(img, mask=img.split()[3])  # Sử dụng kênh alpha làm mask
                return background
            return img
        return None
    except Exception as e:
        st.error(f"Lỗi khi tải ảnh từ URL: {e}")
        return None

def get_download_link(file_content, file_name, display_text):
    """Tạo link tải xuống cho các file được tạo."""
    b64 = base64.b64encode(file_content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}" class="download-button">{display_text}</a>'
    return href

def export_to_pdf(survey_data, images, panel_notes=None):
    """Tạo file PDF từ dữ liệu khảo sát."""
    # Import các module cần thiết
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    buffer = io.BytesIO()
    
    # Cố gắng đăng ký font hỗ trợ tiếng Việt
    try:
        # Tìm kiếm font ở nhiều vị trí
        font_paths = [
            'assets/DejaVuSans.ttf',
            'assets/fonts/DejaVuSans.ttf',
            'DejaVuSans.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'  # Linux path
        ]
        
        font_registered = False
        for path in font_paths:
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', path))
                font_registered = True
                break
            except:
                continue
                
        font_name = 'DejaVuSans' if font_registered else 'Helvetica'
    except:
        font_name = 'Helvetica'
    
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Tạo style cho tiêu đề và nội dung
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=16,
        alignment=1,
        spaceAfter=12
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        spaceAfter=6
    )
    
    content = []
    
    # Tiêu đề báo cáo
    content.append(Paragraph("BÁO CÁO KHẢO SÁT", title_style))
    content.append(Spacer(1, 20))
    
    # Thông tin công ty và khảo sát
    header_data = [
        ["Tên công ty:", survey_data['header']['company_name']],
        ["Địa chỉ:", survey_data['header']['address']],
        ["Số điện thoại:", survey_data['header']['phone']],
        ["Ngày khảo sát:", survey_data['header']['survey_date']],
    ]
    
    t = Table(header_data, colWidths=[120, 300])
    t.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), font_name, 10),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.white),
    ]))
    content.append(t)
    content.append(Spacer(1, 10))
    
    # Người tham gia
    content.append(Paragraph("Thành viên tham gia:", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name)))
    participants = survey_data['header']['participants'].split('\n')
    for p in participants:
        if p.strip():
            content.append(Paragraph(f"• {p}", normal_style))
    content.append(Spacer(1, 10))
    
    # Người khảo sát
    content.append(Paragraph("Người khảo sát:", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name)))
    surveyors = survey_data['header']['surveyors'].split('\n')
    for s in surveyors:
        if s.strip():
            content.append(Paragraph(f"• {s}", normal_style))
    content.append(Spacer(1, 20))
    
    # Chi tiết khảo sát
    content.append(Paragraph("CHI TIẾT KHẢO SÁT", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
    content.append(Spacer(1, 10))
    
    detail_data = [
        ["Khu vực:", survey_data['detail'][0]],
        ["Thiết bị:", survey_data['detail'][1]],
        ["Mô tả tổn thất/thông số kỹ thuật:", survey_data['detail'][2]],
    ]
    
    t = Table(detail_data, colWidths=[150, 270])
    t.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), font_name, 10),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.white),
    ]))
    content.append(t)
    content.append(Spacer(1, 20))
    
    # Hình ảnh
    if images:
        content.append(Paragraph("HÌNH ẢNH KHẢO SÁT", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
        content.append(Spacer(1, 10))
        
        for i, img in enumerate(images):
            if img:
                # Chuyển đổi sang RGB nếu là RGBA
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                
                # Chuyển đổi PIL Image sang bytesIO để ReportLab có thể sử dụng
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG')
                img_byte_arr.seek(0)
                
                # Thêm caption cho hình ảnh
                content.append(Paragraph(f"Hình {i+1}:", normal_style))
                
                # Tính toán kích thước ảnh phù hợp (tối đa 400x300, giữ tỷ lệ)
                img_width, img_height = img.size
                ratio = min(400/img_width, 300/img_height) if img_width > 0 and img_height > 0 else 1
                new_width = img_width * ratio
                new_height = img_height * ratio
                
                # Thêm hình ảnh vào tài liệu PDF
                img_reportlab = RLImage(img_byte_arr, width=new_width, height=new_height)
                content.append(img_reportlab)
                content.append(Spacer(1, 10))
    
    # Panel Notes
    if panel_notes and len(panel_notes) > 0:
        content.append(Paragraph("PANEL NOTES", ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=font_name, alignment=1)))
        content.append(Spacer(1, 10))
        
        for note_idx, note in enumerate(panel_notes):
            content.append(Paragraph(f"Panel Note #{note_idx+1}", ParagraphStyle('Heading3', parent=styles['Heading3'], fontName=font_name)))
            
            note_data = [
                ["Khu vực:", note['area']],
                ["Thiết bị:", note['device']],
                ["Mô tả tổn thất/thông số kỹ thuật:", note['findings']],
            ]
            
            t = Table(note_data, colWidths=[150, 270])
            t.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.white),
            ]))
            content.append(t)
            content.append(Spacer(1, 10))
            
            # Hình ảnh của panel note
            if note.get('images') and len(note['images']) > 0:
                for img_idx, img_url in enumerate(note['images']):
                    img = load_image_from_url(img_url)
                    if img:
                        # Chuyển đổi sang RGB nếu là RGBA
                        if img.mode == 'RGBA':
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[3])
                            img = background
                        
                        # Chuyển đổi PIL Image sang bytesIO để ReportLab có thể sử dụng
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG')
                        img_byte_arr.seek(0)
                        
                        # Thêm caption cho hình ảnh
                        content.append(Paragraph(f"Hình {note_idx+1}.{img_idx+1}:", normal_style))
                        
                        # Tính toán kích thước ảnh phù hợp
                        img_width, img_height = img.size
                        ratio = min(400/img_width, 300/img_height) if img_width > 0 and img_height > 0 else 1
                        new_width = img_width * ratio
                        new_height = img_height * ratio
                        
                        # Thêm hình ảnh vào tài liệu PDF
                        img_reportlab = RLImage(img_byte_arr, width=new_width, height=new_height)
                        content.append(img_reportlab)
                        content.append(Spacer(1, 10))
            
            content.append(Spacer(1, 10))
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    content.append(Spacer(1, 20))
    content.append(Paragraph(f"Báo cáo được xuất ngày: {current_date}", normal_style))
    content.append(Paragraph(f"Người xuất báo cáo: {st.session_state.user.get('full_name', '')}", normal_style))
    
    doc.build(content)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def export_to_word(survey_data, images, panel_notes=None):
    """Tạo file Word từ dữ liệu khảo sát."""
    # Import các module cần thiết
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Thiết lập font và cỡ chữ mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Tiêu đề
    heading = doc.add_heading('BÁO CÁO KHẢO SÁT', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin công ty và khảo sát
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    rows = table.rows
    cells = rows[0].cells
    cells[0].text = 'Tên công ty:'
    cells[1].text = survey_data['header']['company_name']
    
    cells = rows[1].cells
    cells[0].text = 'Địa chỉ:'
    cells[1].text = survey_data['header']['address']
    
    cells = rows[2].cells
    cells[0].text = 'Số điện thoại:'
    cells[1].text = survey_data['header']['phone']
    
    cells = rows[3].cells
    cells[0].text = 'Ngày khảo sát:'
    cells[1].text = survey_data['header']['survey_date']
    
    doc.add_paragraph('')
    
    # Người tham gia
    doc.add_heading('Thành viên tham gia:', level=2)
    participants = survey_data['header']['participants'].split('\n')
    for p in participants:
        if p.strip():
            doc.add_paragraph(f"• {p}", style='List Bullet')
    
    # Người khảo sát
    doc.add_heading('Người khảo sát:', level=2)
    surveyors = survey_data['header']['surveyors'].split('\n')
    for s in surveyors:
        if s.strip():
            doc.add_paragraph(f"• {s}", style='List Bullet')
    
    doc.add_paragraph('')
    
    # Chi tiết khảo sát
    heading = doc.add_heading('CHI TIẾT KHẢO SÁT', level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    rows = table.rows
    cells = rows[0].cells
    cells[0].text = 'Khu vực:'
    cells[1].text = survey_data['detail'][0]
    
    cells = rows[1].cells
    cells[0].text = 'Thiết bị:'
    cells[1].text = survey_data['detail'][1]
    
    cells = rows[2].cells
    cells[0].text = 'Mô tả tổn thất/thông số kỹ thuật:'
    cells[1].text = survey_data['detail'][2]
    
    doc.add_paragraph('')
    
    # Hình ảnh
    if images:
        heading = doc.add_heading('HÌNH ẢNH KHẢO SÁT', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, img in enumerate(images):
            if img:
                # Chuyển đổi sang RGB nếu là RGBA
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                
                # Lưu ảnh tạm thời để chèn vào Word
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG')
                img_byte_arr.seek(0)
                
                # Thêm caption và hình ảnh
                doc.add_paragraph(f"Hình {i+1}:")
                doc.add_picture(img_byte_arr, width=Inches(6))
                doc.add_paragraph('')
    
    # Panel Notes
    if panel_notes and len(panel_notes) > 0:
        heading = doc.add_heading('PANEL NOTES', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for note_idx, note in enumerate(panel_notes):
            doc.add_heading(f"Panel Note #{note_idx+1}", level=3)
            
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            
            rows = table.rows
            cells = rows[0].cells
            cells[0].text = 'Khu vực:'
            cells[1].text = note['area']
            
            cells = rows[1].cells
            cells[0].text = 'Thiết bị:'
            cells[1].text = note['device']
            
            cells = rows[2].cells
            cells[0].text = 'Mô tả tổn thất/thông số kỹ thuật:'
            cells[1].text = note['findings']
            
            doc.add_paragraph('')
            
            # Hình ảnh của panel note
            if note.get('images') and len(note['images']) > 0:
                for img_idx, img_url in enumerate(note['images']):
                    img = load_image_from_url(img_url)
                    if img:
                        # Chuyển đổi sang RGB nếu là RGBA
                        if img.mode == 'RGBA':
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[3])
                            img = background
                        
                        # Lưu ảnh tạm thời để chèn vào Word
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG')
                        img_byte_arr.seek(0)
                        
                        # Thêm caption và hình ảnh
                        doc.add_paragraph(f"Hình {note_idx+1}.{img_idx+1}:")
                        doc.add_picture(img_byte_arr, width=Inches(6))
                        doc.add_paragraph('')
            
            doc.add_paragraph('')
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    doc.add_paragraph('')
    doc.add_paragraph(f"Báo cáo được xuất ngày: {current_date}")
    doc.add_paragraph(f"Người xuất báo cáo: {st.session_state.user.get('full_name', '')}")
    
    # Lưu vào memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    docx = buffer.getvalue()
    buffer.close()
    return docx

# --- Hàm xóa khảo sát ---
def delete_survey_from_supabase(supabase, survey_id):
    """Xóa khảo sát từ Supabase"""
    if not supabase or not survey_id:
        return False
        
    try:
        # Kiểm tra quyền nếu không phải admin
        if st.session_state.user["role"] != "admin":
            response = supabase.table('surveys').select('created_by').eq('id', survey_id).execute()
            if not response.data or response.data[0]["created_by"] != st.session_state.user["id"]:
                st.error("Bạn không có quyền xóa khảo sát này!")
                return False
        
        # Xóa khảo sát (các panel_notes sẽ tự động bị xóa do có constraint ON DELETE CASCADE)
        response = supabase.table('surveys').delete().eq('id', survey_id).execute()
        return True
    except Exception as e:
        st.error(f"Lỗi khi xóa khảo sát: {str(e)}")
        return False

# --- Hàm lấy chi tiết khảo sát ---
def get_survey_detail(supabase, survey_id):
    """Lấy chi tiết khảo sát từ Supabase"""
    if not supabase or not survey_id:
        return None
        
    try:
        # Lấy thông tin khảo sát
        response = supabase.table('surveys').select('*').eq('id', survey_id).execute()
        if not response.data or len(response.data) == 0:
            st.error(f"Không tìm thấy khảo sát ID: {survey_id}")
            return None
            
        survey = response.data[0]
        
        # Lấy danh sách panel notes
        panel_notes_response = supabase.table('panel_notes').select('*').eq('survey_id', survey_id).execute()
        panel_notes = panel_notes_response.data if panel_notes_response.data else []
        
        # Tạo đối tượng dữ liệu đầy đủ
        survey_data = {
            'header': {
                'company_name': survey['company_name'],
                'address': survey['address'],
                'phone': survey['phone'],
                'survey_date': survey['survey_date'],
                'participants': survey['participants'],
                'surveyors': survey['surveyors']
            },
            'detail': [
                survey['area'],
                survey['device'],
                survey['findings']
            ],
            'image_urls': survey['images'] if survey['images'] else [],
            'panel_notes': panel_notes
        }
        
        return survey_data
    except Exception as e:
        st.error(f"Lỗi khi lấy chi tiết khảo sát: {str(e)}")
        st.error(traceback.format_exc())
        return None

# Giao diện chính
st.title("📋 Xem danh sách khảo sát")

# Kết nối Supabase
supabase = init_supabase()

if not supabase:
    st.error("😢 Không thể kết nối đến cơ sở dữ liệu Supabase.")
    st.warning("Vui lòng cấu hình kết nối Supabase trong Streamlit Secrets.")
    st.stop()

# Thêm nút tạo khảo sát mới
if st.button("➕ Tạo khảo sát mới"):
    # Xóa ID khảo sát đang chỉnh sửa nếu có
    if 'editing_survey_id' in st.session_state:
        st.session_state.editing_survey_id = None
    # Chuyển đến trang nhập liệu
    try:
        st.switch_page("../Getnotes_Onsite.py")
    except Exception:
        try:
            st.switch_page("Getnotes_Onsite.py")
        except Exception:
            st.error("Không thể chuyển đến trang nhập liệu. Vui lòng quay lại trang chủ.")

# Tạo tab cho các chức năng
tab1, tab2 = st.tabs(["Danh sách khảo sát", "Tìm kiếm"])

with tab1:
    st.subheader("Danh sách khảo sát")
    
    # Filter theo người tạo
    view_options = ["Tất cả khảo sát", "Khảo sát của tôi"]
    if st.session_state.user["role"] == "admin":
        selected_view = st.radio("Hiển thị:", view_options)
    else:
        selected_view = "Khảo sát của tôi"
    
    # Lấy danh sách khảo sát từ Supabase
    try:
        if selected_view == "Tất cả khảo sát" and st.session_state.user["role"] == "admin":
            # Admin xem tất cả các khảo sát
            response = supabase.table('surveys').select('*, users!inner(full_name)').order('created_at', desc=True).execute()
        else:
            # Người dùng thông thường chỉ xem các khảo sát của mình
            response = supabase.table('surveys').select('*, users!inner(full_name)').eq('created_by', st.session_state.user["id"]).order('created_at', desc=True).execute()
        
        if response.data and len(response.data) > 0:
            st.write(f"Tìm thấy {len(response.data)} khảo sát")
            
            # Hiển thị dưới dạng bảng
            surveys = []
            for survey in response.data:
                created_at = datetime.datetime.fromisoformat(survey['created_at'].replace('Z', '+00:00'))
                formatted_date = created_at.strftime("%d/%m/%Y %H:%M")
                
                surveys.append({
                    "ID": survey['id'],
                    "Công ty": survey['company_name'],
                    "Ngày khảo sát": survey['survey_date'],
                    "Người tạo": survey['users']['full_name'],
                    "Thời gian tạo": formatted_date
                })
            
            # Hiển thị bảng khảo sát
            st.dataframe(surveys, use_container_width=True)
            
            # Chọn khảo sát để xem chi tiết
            survey_ids = [s['id'] for s in response.data]
            survey_names = [f"{s['company_name']} ({s['survey_date']})" for s in response.data]
            options = dict(zip(survey_ids, survey_names))
            
            selected_survey = st.selectbox("Chọn khảo sát để xem chi tiết:", survey_ids, format_func=lambda x: options[x])
            
            if selected_survey:
                st.subheader("Thao tác")
                
                # Nút xem chi tiết, chỉnh sửa và xóa
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("📄 Xem chi tiết", key="view_btn"):
                        st.session_state.selected_survey_id = selected_survey
                        
                with col2:
                    if st.button("✏️ Chỉnh sửa", key="edit_btn"):
                        # Kiểm tra quyền nếu không phải admin
                        if st.session_state.user["role"] != "admin":
                            selected_survey_data = [s for s in response.data if s['id'] == selected_survey][0]
                            if selected_survey_data["created_by"] != st.session_state.user["id"]:
                                st.error("Bạn không có quyền chỉnh sửa khảo sát này!")
                                st.stop()
                        
                        # Thiết lập ID khảo sát để chỉnh sửa
                        st.session_state.editing_survey_id = selected_survey
                        # Chuyển đến trang nhập liệu
                        try:
                            st.switch_page("../Getnotes_Onsite.py")
                        except Exception:
                            try:
                                st.switch_page("Getnotes_Onsite.py")
                            except Exception:
                                st.error("Không thể chuyển đến trang nhập liệu. Vui lòng quay lại trang chủ.")
                        
                with col3:
                    if st.button("🗑️ Xóa", key="delete_btn"):
                        # Tạo hộp thoại xác nhận
                        st.session_state.confirm_delete = selected_survey
                
                # Xác nhận xóa
                if 'confirm_delete' in st.session_state and st.session_state.confirm_delete:
                    confirm_col1, confirm_col2 = st.columns(2)
                    with confirm_col1:
                        if st.button("✓ Xác nhận xóa", key="confirm_delete_btn"):
                            if delete_survey_from_supabase(supabase, st.session_state.confirm_delete):
                                st.success(f"Đã xóa khảo sát ID: {st.session_state.confirm_delete}")
                                st.session_state.confirm_delete = None
                                st.rerun()
                            else:
                                st.error("Không thể xóa khảo sát")
                                
                    with confirm_col2:
                        if st.button("✗ Hủy", key="cancel_delete_btn"):
                            st.session_state.confirm_delete = None
                            st.rerun()
                
                # Hiển thị chi tiết khảo sát nếu đã chọn
                if 'selected_survey_id' in st.session_state and st.session_state.selected_survey_id:
                    st.subheader("Chi tiết khảo sát")
                    
                    # Lấy dữ liệu khảo sát
                    survey_data = get_survey_detail(supabase, st.session_state.selected_survey_id)
                    
                    if survey_data:
                        # Hiển thị thông tin
                        st.write("### Thông tin công ty")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write(f"**Tên công ty:** {survey_data['header']['company_name']}")
                            st.write(f"**Địa chỉ:** {survey_data['header']['address']}")
                        with col2:
                            st.write(f"**Số điện thoại:** {survey_data['header']['phone']}")
                            st.write(f"**Ngày khảo sát:** {survey_data['header']['survey_date']}")
                        
                        st.write("### Thành viên tham gia")
                        participants = survey_data['header']['participants'].split('\n')
                        for p in participants:
                            if p.strip():
                                st.write(f"• {p}")
                        
                        st.write("### Người khảo sát")
                        surveyors = survey_data['header']['surveyors'].split('\n')
                        for s in surveyors:
                            if s.strip():
                                st.write(f"• {s}")
                        
                        st.write("### Chi tiết khảo sát")
                        st.write(f"**Khu vực:** {survey_data['detail'][0]}")
                        st.write(f"**Thiết bị:** {survey_data['detail'][1]}")
                        st.write(f"**Mô tả tổn thất/thông số kỹ thuật:** {survey_data['detail'][2]}")
                        
                        # Hiển thị hình ảnh
                        if survey_data['image_urls'] and len(survey_data['image_urls']) > 0:
                            st.write("### Hình ảnh khảo sát")
                            img_cols = st.columns(min(3, len(survey_data['image_urls'])))
                            for idx, img_url in enumerate(survey_data['image_urls']):
                                with img_cols[idx % 3]:
                                    st.image(img_url, caption=f"Hình {idx+1}", width=200)
                        
                        # Hiển thị panel notes
                        if survey_data['panel_notes'] and len(survey_data['panel_notes']) > 0:
                            st.write("### Panel Notes")
                            
                            for idx, note in enumerate(survey_data['panel_notes']):
                                with st.expander(f"Panel Note #{idx+1}", expanded=False):
                                    st.write(f"**Khu vực:** {note['area']}")
                                    st.write(f"**Thiết bị:** {note['device']}")
                                    st.write(f"**Mô tả tổn thất/thông số kỹ thuật:** {note['findings']}")
                                    
                                    # Hiển thị hình ảnh của panel note
                                    if note.get('images') and len(note['images']) > 0:
                                        st.write("#### Hình ảnh")
                                        note_img_cols = st.columns(min(3, len(note['images'])))
                                        for img_idx, img_url in enumerate(note['images']):
                                            with note_img_cols[img_idx % 3]:
                                                st.image(img_url, caption=f"Hình {img_idx+1}", width=200)
                        
                        # Nút xuất báo cáo
                        st.subheader("Xuất báo cáo")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if st.button("📄 Xuất file PDF"):
                                with st.spinner("Đang tạo file PDF..."):
                                    # Tải hình ảnh từ URL
                                    images = []
                                    for url in survey_data['image_urls']:
                                        img = load_image_from_url(url)
                                        if img:
                                            images.append(img)
                                    
                                    # Tạo file PDF
                                    pdf_data = export_to_pdf(survey_data, images, survey_data['panel_notes'])
                                    
                                    # Tạo tên file
                                    company_name_safe = survey_data['header']['company_name'].replace(' ', '_')
                                    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"bao_cao_khao_sat_{company_name_safe}_{date_str}.pdf"
                                    
                                    # Hiển thị link tải xuống
                                    st.markdown(
                                        get_download_link(pdf_data, filename, "📥 Tải xuống file PDF"),
                                        unsafe_allow_html=True
                                    )
                        
                        with col2:
                            if st.button("📄 Xuất file Word"):
                                with st.spinner("Đang tạo file Word..."):
                                    # Tải hình ảnh từ URL
                                    images = []
                                    for url in survey_data['image_urls']:
                                        img = load_image_from_url(url)
                                        if img:
                                            images.append(img)
                                    
                                    # Tạo file Word
                                    docx_data = export_to_word(survey_data, images, survey_data['panel_notes'])
                                    
                                    # Tạo tên file
                                    company_name_safe = survey_data['header']['company_name'].replace(' ', '_')
                                    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"bao_cao_khao_sat_{company_name_safe}_{date_str}.docx"
                                    
                                    # Hiển thị link tải xuống
                                    st.markdown(
                                        get_download_link(docx_data, filename, "📥 Tải xuống file Word"),
                                        unsafe_allow_html=True
                                    )
                    else:
                        st.error("Không thể tải dữ liệu khảo sát")
        else:
            st.info("Không có khảo sát nào. Nhấn nút 'Tạo khảo sát mới' để bắt đầu.")
    except Exception as e:
        st.error(f"Lỗi khi tải danh sách khảo sát: {str(e)}")
        st.error(traceback.format_exc())

with tab2:
    st.subheader("Tìm kiếm khảo sát")
    
    search_term = st.text_input("Nhập từ khóa tìm kiếm:")
    search_type = st.radio("Tìm kiếm theo:", ["Tên công ty", "Khu vực", "Thiết bị", "Mô tả"])
    
    if st.button("🔍 Tìm kiếm"):
        if search_term:
            try:
                # Xác định trường tìm kiếm
                search_field = "company_name"
                if search_type == "Khu vực":
                    search_field = "area"
                elif search_type == "Thiết bị":
                    search_field = "device"
                elif search_type == "Mô tả":
                    search_field = "findings"
                
                # Thực hiện tìm kiếm
                if st.session_state.user["role"] == "admin":
                    # Admin có thể tìm kiếm tất cả các khảo sát
                    response = supabase.table('surveys').select('*, users!inner(full_name)').ilike(search_field, f"%{search_term}%").order('created_at', desc=True).execute()
                else:
                    # Người dùng thông thường chỉ tìm kiếm trong các khảo sát của mình
                    response = supabase.table('surveys').select('*, users!inner(full_name)').eq('created_by', st.session_state.user["id"]).ilike(search_field, f"%{search_term}%").order('created_at', desc=True).execute()
                
                if response.data and len(response.data) > 0:
                    st.write(f"Tìm thấy {len(response.data)} kết quả")
                    
                    # Hiển thị dưới dạng bảng
                    search_results = []
                    for survey in response.data:
                        created_at = datetime.datetime.fromisoformat(survey['created_at'].replace('Z', '+00:00'))
                        formatted_date = created_at.strftime("%d/%m/%Y %H:%M")
                        
                        search_results.append({
                            "ID": survey['id'],
                            "Công ty": survey['company_name'],
                            "Khu vực": survey['area'],
                            "Thiết bị": survey['device'],
                            "Ngày khảo sát": survey['survey_date'],
                            "Người tạo": survey['users']['full_name'],
                            "Thời gian tạo": formatted_date
                        })
                    
                    # Hiển thị bảng kết quả tìm kiếm
                    st.dataframe(search_results, use_container_width=True)
                    
                    # Chọn khảo sát để xem chi tiết
                    survey_ids = [s['id'] for s in response.data]
                    survey_names = [f"{s['company_name']} ({s['survey_date']})" for s in response.data]
                    options = dict(zip(survey_ids, survey_names))
                    
                    selected_survey = st.selectbox("Chọn khảo sát để xem chi tiết:", survey_ids, format_func=lambda x: options[x], key="search_select")
                    
                    if selected_survey:
                        if st.button("📄 Xem chi tiết", key="search_view_btn"):
                            st.session_state.selected_survey_id = selected_survey
                            st.rerun()
                else:
                    st.info(f"Không tìm thấy kết quả nào cho từ khóa '{search_term}'")
            except Exception as e:
                st.error(f"Lỗi khi tìm kiếm: {str(e)}")
                st.error(traceback.format_exc())